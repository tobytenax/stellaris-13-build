"""
Stellaris-13 Celestial Blueprint Generator

Generates professional multi-chapter natal chart documents in DOCX format.
Each chapter is AI-generated based on the computed chart data.
"""

import os
import logging
from datetime import datetime
from typing import Callable, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from blueprint_prompts import (
    ASTROLOGER_PERSONA,
    get_title_page_content,
    get_chapter_prompts,
    build_chart_summary
)

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT STYLING
# ═══════════════════════════════════════════════════════════════════════════════

def setup_document_styles(doc: Document):
    """Configure document styles for the blueprint."""
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(11)
    
    # Title style
    title_style = doc.styles.add_style('BlueprintTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Georgia'
    title_style.font.size = Pt(36)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(45, 19, 44)  # Deep purple
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(6)
    
    # Subtitle style
    subtitle_style = doc.styles.add_style('BlueprintSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_style.font.name = 'Georgia'
    subtitle_style.font.size = Pt(18)
    subtitle_style.font.italic = True
    subtitle_style.font.color.rgb = RGBColor(128, 19, 54)
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_after = Pt(24)
    
    # Chapter heading style
    chapter_style = doc.styles.add_style('ChapterHeading', WD_STYLE_TYPE.PARAGRAPH)
    chapter_style.font.name = 'Georgia'
    chapter_style.font.size = Pt(24)
    chapter_style.font.bold = True
    chapter_style.font.color.rgb = RGBColor(45, 19, 44)
    chapter_style.paragraph_format.space_before = Pt(36)
    chapter_style.paragraph_format.space_after = Pt(18)
    chapter_style.paragraph_format.page_break_before = True
    
    # Body text style  
    body_style = doc.styles.add_style('BlueprintBody', WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.name = 'Georgia'
    body_style.font.size = Pt(11)
    body_style.paragraph_format.space_after = Pt(12)
    body_style.paragraph_format.line_spacing = 1.5
    
    # Quote/emphasis style
    quote_style = doc.styles.add_style('BlueprintQuote', WD_STYLE_TYPE.PARAGRAPH)
    quote_style.font.name = 'Georgia'
    quote_style.font.size = Pt(14)
    quote_style.font.italic = True
    quote_style.font.color.rgb = RGBColor(128, 19, 54)
    quote_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote_style.paragraph_format.space_before = Pt(24)
    quote_style.paragraph_format.space_after = Pt(24)
    
    return doc


def add_page_break(doc: Document):
    """Add a page break."""
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════

def create_title_page(doc: Document, chart: dict, birth_data: dict):
    """Create the title page of the blueprint."""
    
    title_content = get_title_page_content(chart)
    name = title_content['name']
    
    # Add spacing at top
    for _ in range(3):
        doc.add_paragraph()
    
    # Main title
    title = doc.add_paragraph("THE COMPLETE CELESTIAL BLUEPRINT", style='BlueprintTitle')
    
    # Subtitle
    subtitle = doc.add_paragraph("IAU 13-SIGN ASTRONOMICAL EDITION", style='BlueprintSubtitle')
    
    # Name
    doc.add_paragraph()
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(f"An Exhaustive Astrological Reading for {name.upper()}")
    name_run.font.name = 'Georgia'
    name_run.font.size = Pt(16)
    name_run.font.italic = True
    
    # Birth data
    birth_str = format_birth_data(birth_data)
    birth_para = doc.add_paragraph()
    birth_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    birth_run = birth_para.add_run(birth_str)
    birth_run.font.name = 'Georgia'
    birth_run.font.size = Pt(12)
    
    # Big symbol
    doc.add_paragraph()
    doc.add_paragraph()
    symbol_para = doc.add_paragraph()
    symbol_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Use Ophiuchus symbol if they have Ophiuchus placements, otherwise their rising sign symbol
    placements = chart.get('placements', {})
    has_ophiuchus = any(p.get('iau_constellation') == 'Ophiuchus' for p in placements.values())
    
    if has_ophiuchus:
        symbol = "⛎"
        symbol_label = "OPHIUCHUS — THE SERPENT-BEARER"
    else:
        rising = title_content['rising_sign']
        symbol = get_sign_symbol(rising)
        symbol_label = f"{rising.upper()} RISING"
    
    symbol_run = symbol_para.add_run(symbol)
    symbol_run.font.size = Pt(72)
    
    label_para = doc.add_paragraph()
    label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = label_para.add_run(symbol_label)
    label_run.font.name = 'Georgia'
    label_run.font.size = Pt(14)
    label_run.font.bold = True
    label_run.font.color.rgb = RGBColor(128, 19, 54)
    
    # Core signature
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig = f"Sun in {title_content['sun_sign']} • Moon in {title_content['moon_sign']} • {title_content['rising_sign']} Rising"
    sig_run = sig_para.add_run(sig)
    sig_run.font.name = 'Georgia'
    sig_run.font.size = Pt(12)
    sig_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("Computed using Stellaris-13")
    footer_run.font.name = 'Georgia'
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer2_run = footer2.add_run("IAU Constellation Boundaries (1930) • Swiss Ephemeris • Whole Sign Houses")
    footer2_run.font.name = 'Georgia'
    footer2_run.font.size = Pt(10)
    footer2_run.font.color.rgb = RGBColor(128, 128, 128)
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated {datetime.now().strftime('%B %d, %Y')}")
    date_run.font.name = 'Georgia'
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(128, 128, 128)


def format_birth_data(birth_data: dict) -> str:
    """Format birth data for display."""
    year = birth_data.get('year', '?')
    month = birth_data.get('month', '?')
    day = birth_data.get('day', '?')
    hour = birth_data.get('hour', '?')
    minute = birth_data.get('minute', '?')
    
    # Format date
    try:
        from datetime import date
        d = date(year, month, day)
        date_str = d.strftime('%B %d, %Y')
    except:
        date_str = f"{month}/{day}/{year}"
    
    # Format time
    try:
        h = int(hour)
        m = int(minute)
        am_pm = "AM" if h < 12 else "PM"
        h_12 = h if h <= 12 else h - 12
        if h_12 == 0:
            h_12 = 12
        time_str = f"{h_12}:{m:02d} {am_pm}"
    except:
        time_str = f"{hour}:{minute}"
    
    lat = birth_data.get('lat', 0)
    lon = birth_data.get('lon', 0)
    
    return f"Born {date_str} — {time_str} — ({lat:.2f}°, {lon:.2f}°)"


def get_sign_symbol(sign: str) -> str:
    """Get the Unicode symbol for a zodiac sign."""
    symbols = {
        'Aries': '♈', 'Taurus': '♉', 'Gemini': '♊', 'Cancer': '♋',
        'Leo': '♌', 'Virgo': '♍', 'Libra': '♎', 'Scorpius': '♏',
        'Scorpio': '♏', 'Sagittarius': '♐', 'Capricorn': '♑', 
        'Capricornus': '♑', 'Aquarius': '♒', 'Pisces': '♓',
        'Ophiuchus': '⛎'
    }
    return symbols.get(sign, '✦')


# ═══════════════════════════════════════════════════════════════════════════════
# CHART DATA TABLE
# ═══════════════════════════════════════════════════════════════════════════════

def add_positions_table(doc: Document, chart: dict):
    """Add a table showing all planetary positions."""
    
    placements = chart['placements']
    angles = chart['angles']
    
    # Chapter heading (but without page break for this one)
    heading = doc.add_paragraph("Natal Positions Reference")
    heading.style = doc.styles['ChapterHeading']
    heading.paragraph_format.page_break_before = False
    
    # Create table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    headers = ['Body', 'IAU Constellation', 'Degree', 'House', 'Tropical']
    for i, text in enumerate(headers):
        header_cells[i].text = text
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    
    # Planet rows
    bodies = ['Sun', 'Moon', 'Mercury', 'Venus', 'Mars', 'Jupiter', 'Saturn', 'Uranus', 'Neptune', 'Pluto',
              'North Node', 'South Node', 'Lilith']
    
    for body in bodies:
        if body in placements:
            p = placements[body]
            row = table.add_row().cells
            row[0].text = body
            row[1].text = p.get('iau_constellation', '?')
            row[2].text = f"{p.get('iau_degree', 0):.2f}°"
            row[3].text = f"H{p.get('house', '?')}"
            row[4].text = p.get('tropical_sign', '?')
            
            # Highlight Ophiuchus
            if p.get('iau_constellation') == 'Ophiuchus':
                row[1].paragraphs[0].runs[0].font.bold = True
                row[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(128, 19, 54)
    
    # Angles section
    doc.add_paragraph()
    angles_heading = doc.add_paragraph("Angles")
    angles_heading.runs[0].font.bold = True
    angles_heading.runs[0].font.size = Pt(12)
    
    angles_table = doc.add_table(rows=1, cols=4)
    angles_table.style = 'Table Grid'
    
    header_cells = angles_table.rows[0].cells
    for i, text in enumerate(['Angle', 'IAU Constellation', 'IAU Degree', 'Tropical']):
        header_cells[i].text = text
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    
    for angle in ['Ascendant', 'MC', 'Descendant', 'IC']:
        if angle in angles:
            a = angles[angle]
            row = angles_table.add_row().cells
            row[0].text = angle
            row[1].text = a.get('iau_constellation', '?')
            row[2].text = f"{a.get('iau_degree', 0):.2f}°"
            row[3].text = f"{a.get('tropical_sign', '?')} {a.get('tropical_degree', 0):.2f}°"


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER GENERATION
# ═══════════════════════════════════════════════════════════════════════════════

def add_chapter(doc: Document, title: str, content: str):
    """Add a chapter to the document."""
    
    # Chapter heading
    heading = doc.add_paragraph(title, style='ChapterHeading')
    
    # Split content into paragraphs and add each
    paragraphs = content.strip().split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            # Clean up the text
            clean_text = para_text.strip().replace('\n', ' ')
            para = doc.add_paragraph(clean_text, style='BlueprintBody')


def add_closing_page(doc: Document, chart: dict):
    """Add the closing page with symbols and attribution."""
    
    doc.add_page_break()
    
    # Get dominant signs for the symbol display
    placements = chart.get('placements', {})
    sun = placements.get('Sun', {}).get('iau_constellation', 'Unknown')
    moon = placements.get('Moon', {}).get('iau_constellation', 'Unknown')
    rising = chart.get('angles', {}).get('Ascendant', {}).get('iau_constellation', 'Unknown')
    
    for _ in range(6):
        doc.add_paragraph()
    
    # Symbols
    symbols = f"{get_sign_symbol(sun)} {get_sign_symbol(moon)} {get_sign_symbol(rising)}"
    sym_para = doc.add_paragraph()
    sym_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sym_run = sym_para.add_run(symbols)
    sym_run.font.size = Pt(48)
    
    # Divider
    doc.add_paragraph()
    div_para = doc.add_paragraph()
    div_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_run = div_para.add_run("─" * 30)
    div_run.font.color.rgb = RGBColor(45, 19, 44)
    
    # Attribution
    doc.add_paragraph()
    attr1 = doc.add_paragraph()
    attr1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    attr1_run = attr1.add_run("Generated by Stellaris-13")
    attr1_run.font.name = 'Georgia'
    attr1_run.font.size = Pt(10)
    attr1_run.font.italic = True
    attr1_run.font.color.rgb = RGBColor(128, 128, 128)
    
    attr2 = doc.add_paragraph()
    attr2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    attr2_run = attr2.add_run("IAU Constellation Boundaries (1930) • Swiss Ephemeris • Whole Sign Houses")
    attr2_run.font.name = 'Georgia'
    attr2_run.font.size = Pt(10)
    attr2_run.font.italic = True
    attr2_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Quote
    doc.add_paragraph()
    doc.add_paragraph()
    quote_para = doc.add_paragraph()
    quote_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote_run = quote_para.add_run('"The stars incline, they do not compel."')
    quote_run.font.name = 'Georgia'
    quote_run.font.size = Pt(12)
    quote_run.font.italic = True


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN GENERATOR
# ═══════════════════════════════════════════════════════════════════════════════

def generate_blueprint(
    chart: dict,
    birth_data: dict,
    ai_caller: Callable[[str], str],
    output_path: str,
    progress_callback: Optional[Callable[[int, str], None]] = None
) -> str:
    """
    Generate a complete celestial blueprint document.
    
    Args:
        chart: The computed chart data from engine.compute_chart()
        birth_data: Dict with year, month, day, hour, minute, lat, lon
        ai_caller: Function that takes a prompt string and returns AI response
        output_path: Path to save the .docx file
        progress_callback: Optional callback(step_number, step_description) for progress updates
    
    Returns:
        Path to the generated document
    """
    
    def progress(step: int, desc: str):
        if progress_callback:
            progress_callback(step, desc)
        logger.info(f"Blueprint generation step {step}: {desc}")
    
    progress(0, "Initializing document...")
    
    # Create document
    doc = Document()
    setup_document_styles(doc)
    
    # Title page
    progress(1, "Creating title page...")
    create_title_page(doc, chart, birth_data)
    
    # Positions table
    progress(2, "Adding positions reference table...")
    add_page_break(doc)
    add_positions_table(doc, chart)
    
    # Get chapter prompts
    chapters = get_chapter_prompts(chart)
    total_chapters = len(chapters)
    
    # Generate each chapter
    for i, (title, prompt) in enumerate(chapters):
        step_num = 3 + i
        progress(step_num, f"Writing {title}...")
        
        try:
            # Build the full prompt with persona
            full_prompt = f"{ASTROLOGER_PERSONA}\n\n{prompt}"
            
            # Call AI
            content = ai_caller(full_prompt)
            
            # Add to document
            add_chapter(doc, title, content)
            
        except Exception as e:
            logger.error(f"Error generating {title}: {e}")
            add_chapter(doc, title, f"[Content generation failed: {e}]")
    
    # Closing page
    progress(3 + total_chapters, "Adding closing page...")
    add_closing_page(doc, chart)
    
    # Save document
    progress(4 + total_chapters, "Saving document...")
    doc.save(output_path)
    
    progress(5 + total_chapters, "Blueprint complete!")
    
    return output_path


# ═══════════════════════════════════════════════════════════════════════════════
# TESTING
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    # Test with mock data
    print("Blueprint generator module loaded successfully.")
    print("Use generate_blueprint() to create documents.")
