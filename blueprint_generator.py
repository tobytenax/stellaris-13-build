"""
Stellaris-13 Celestial Blueprint Generator v2
===============================================
Generates professional multi-chapter natal chart documents in DOCX format.
Each chapter is AI-generated based on computed chart data.

Matches the depth and structure of "The Complete Celestial Blueprint" —
9 chapters covering natal chart, transits, progressions, solar arcs,
synastry, past lives, psychology, predictive timing, and special topics.

Usage:
    from blueprint_generator import generate_blueprint
    from engine import compute_chart
    
    chart = compute_chart(year=1985, month=12, day=12, ...)
    
    def my_ai_caller(prompt):
        # Call your AI API here
        return response_text
    
    generate_blueprint(
        chart=chart,
        birth_data={'date': 'December 12, 1985', 'time': '10:47 AM CST', ...},
        ai_caller=my_ai_caller,
        output_path='blueprint.docx'
    )
"""

import os
import re
import logging
from datetime import datetime
from typing import Callable, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from blueprint_prompts import (
    ASTROLOGER_PERSONA,
    get_title_page_content,
    get_chapter_prompts,
    build_chart_summary,
    build_aspect_table
)

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT STYLING
# ═══════════════════════════════════════════════════════════════════════════════

# Color palette — deep celestial theme
VOID = RGBColor(0x0D, 0x0B, 0x1E)       # Near-black blue
GOLD = RGBColor(0xC9, 0xA8, 0x3B)       # Warm gold
STAR_BRIGHT = RGBColor(0xE8, 0xE0, 0xD0) # Warm white
STAR_MID = RGBColor(0xA0, 0x96, 0x86)    # Muted cream
DUST = RGBColor(0x3A, 0x34, 0x4F)        # Deep purple-grey
CHAPTER_BLUE = RGBColor(0x1A, 0x1A, 0x2E) # Dark navy
ACCENT = RGBColor(0x6C, 0x5C, 0xE7)      # Soft purple


def setup_document_styles(doc: Document):
    """Configure document styles for the celestial blueprint."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)
    
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.3
    
    # Title style
    if 'BlueprintTitle' not in [s.name for s in doc.styles]:
        title_style = doc.styles.add_style('BlueprintTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Georgia'
        title_style.font.size = Pt(28)
        title_style.font.bold = True
        title_style.font.color.rgb = CHAPTER_BLUE
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(4)
    
    # Subtitle style
    if 'BlueprintSubtitle' not in [s.name for s in doc.styles]:
        sub_style = doc.styles.add_style('BlueprintSubtitle', WD_STYLE_TYPE.PARAGRAPH)
        sub_style.font.name = 'Georgia'
        sub_style.font.size = Pt(14)
        sub_style.font.italic = True
        sub_style.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        sub_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_style.paragraph_format.space_after = Pt(2)
    
    # Chapter heading style
    if 'ChapterHeading' not in [s.name for s in doc.styles]:
        ch_style = doc.styles.add_style('ChapterHeading', WD_STYLE_TYPE.PARAGRAPH)
        ch_style.font.name = 'Georgia'
        ch_style.font.size = Pt(22)
        ch_style.font.bold = True
        ch_style.font.color.rgb = CHAPTER_BLUE
        ch_style.paragraph_format.space_before = Pt(36)
        ch_style.paragraph_format.space_after = Pt(18)
        ch_style.paragraph_format.page_break_before = True
    
    # Section heading style (## within chapters)
    if 'SectionHeading' not in [s.name for s in doc.styles]:
        sec_style = doc.styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
        sec_style.font.name = 'Georgia'
        sec_style.font.size = Pt(15)
        sec_style.font.bold = True
        sec_style.font.color.rgb = RGBColor(0x44, 0x3A, 0x6E)
        sec_style.paragraph_format.space_before = Pt(24)
        sec_style.paragraph_format.space_after = Pt(10)
    
    # Sub-section heading style (### within sections)
    if 'SubSectionHeading' not in [s.name for s in doc.styles]:
        subsec_style = doc.styles.add_style('SubSectionHeading', WD_STYLE_TYPE.PARAGRAPH)
        subsec_style.font.name = 'Georgia'
        subsec_style.font.size = Pt(13)
        subsec_style.font.bold = True
        subsec_style.font.italic = True
        subsec_style.font.color.rgb = RGBColor(0x55, 0x4D, 0x78)
        subsec_style.paragraph_format.space_before = Pt(16)
        subsec_style.paragraph_format.space_after = Pt(8)
    
    # Epigraph / quote style
    if 'Epigraph' not in [s.name for s in doc.styles]:
        epi_style = doc.styles.add_style('Epigraph', WD_STYLE_TYPE.PARAGRAPH)
        epi_style.font.name = 'Georgia'
        epi_style.font.size = Pt(12)
        epi_style.font.italic = True
        epi_style.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        epi_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        epi_style.paragraph_format.space_before = Pt(24)
        epi_style.paragraph_format.space_after = Pt(24)
        epi_style.paragraph_format.left_indent = Inches(1)
        epi_style.paragraph_format.right_indent = Inches(1)
    
    # Bold emphasis run style (for inline emphasis)
    if 'StrongEmphasis' not in [s.name for s in doc.styles]:
        strong = doc.styles.add_style('StrongEmphasis', WD_STYLE_TYPE.CHARACTER)
        strong.font.bold = True
        strong.font.color.rgb = CHAPTER_BLUE


def add_page_break(doc: Document):
    """Add a page break."""
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════

def create_title_page(doc: Document, chart: dict, birth_data: dict):
    """Create the document title page."""
    tp = get_title_page_content(chart, birth_data)
    
    # Spacer
    for _ in range(4):
        doc.add_paragraph('')
    
    # Main title
    p = doc.add_paragraph(style='BlueprintTitle')
    run = p.add_run('THE COMPLETE CELESTIAL BLUEPRINT')
    
    # Divider
    p = doc.add_paragraph(style='BlueprintSubtitle')
    p.add_run('⛎')
    
    # Subject line
    p = doc.add_paragraph(style='BlueprintSubtitle')
    p.add_run(f'An Exhaustive Astrological Reading for {tp["name"]}')
    
    # Birth data
    p = doc.add_paragraph(style='BlueprintSubtitle')
    p.add_run(f'Born {tp["date"]} — {tp["time"]}')
    
    p = doc.add_paragraph(style='BlueprintSubtitle')
    p.add_run(f'{tp["location"]}')
    
    # Spacer
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # Big three
    p = doc.add_paragraph(style='BlueprintSubtitle')
    p.add_run(f'{tp["sun_sign"]} Sun  ·  {tp["moon_sign"]} Moon  ·  {tp["rising_sign"]} Rising')
    p.paragraph_format.space_before = Pt(24)
    
    # Sect
    p = doc.add_paragraph(style='BlueprintSubtitle')
    p.add_run(f'{tp["sect"]} Chart')
    
    # Spacer
    for _ in range(4):
        doc.add_paragraph('')
    
    # Epigraph
    p = doc.add_paragraph(style='Epigraph')
    p.add_run('"The stars are not pulled to you; you are woven into them."')
    
    # System note
    p = doc.add_paragraph(style='BlueprintSubtitle')
    p.add_run('IAU Astronomical Constellation Boundaries · Swiss Ephemeris · Whole Sign Houses')
    p.paragraph_format.space_before = Pt(36)
    
    p = doc.add_paragraph(style='BlueprintSubtitle')
    p.add_run(f'Computed by Stellaris-13 · {datetime.now().strftime("%B %d, %Y")}')


# ═══════════════════════════════════════════════════════════════════════════════
# POSITIONS TABLE
# ═══════════════════════════════════════════════════════════════════════════════

def add_positions_table(doc: Document, chart: dict):
    """Add a comprehensive positions reference table."""
    p = doc.add_paragraph(style='ChapterHeading')
    p.add_run('Planetary Positions Reference')
    p.paragraph_format.page_break_before = True
    
    placements = chart.get('placements', {})
    angles = chart.get('angles', {})
    
    # Planets table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Body', 'IAU Constellation', 'Degree', 'House', 'Tropical Sign', 'Notes']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    
    body_order = [
        'Sun', 'Moon', 'Mercury', 'Venus', 'Mars', 'Jupiter', 'Saturn',
        'Uranus', 'Neptune', 'Pluto', 'North Node', 'South Node',
        'Chiron', 'Lilith', 'Ceres', 'Pallas', 'Juno', 'Vesta',
        'Vertex', 'Part of Fortune', 'Part of Spirit',
        'Eros', 'Psyche', 'Nessus', 'Pholus', 'Sedna', 'Eris',
        'Orcus', 'Quaoar',
        'Cupido', 'Hades', 'Zeus', 'Kronos', 'Apollon', 
        'Admetos', 'Vulkanus', 'Poseidon'
    ]
    
    for body in body_order:
        p = placements.get(body)
        if p:
            row = table.add_row()
            iau = p.get('iau_constellation', p.get('standard_constellation', '?'))
            deg = p.get('iau_degree', p.get('standard_degree', 0))
            std = p.get('standard_constellation', '?')
            house = str(p.get('house', '?'))
            retro = ' ℞' if p.get('retrograde') else ''
            notes = retro.strip()
            
            row.cells[0].text = body
            row.cells[1].text = iau
            row.cells[2].text = f'{deg:.2f}°'
            row.cells[3].text = f'H{house}'
            row.cells[4].text = std
            row.cells[5].text = notes
            
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
    
    # Angles table
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Angles')
    run.bold = True
    run.font.size = Pt(13)
    
    angle_table = doc.add_table(rows=1, cols=4)
    angle_table.style = 'Light Shading Accent 1'
    angle_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, h in enumerate(['Angle', 'IAU Constellation', 'Degree', 'Tropical']):
        cell = angle_table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    
    for angle_name in ['Ascendant', 'MC', 'Descendant', 'IC']:
        a = angles.get(angle_name, {})
        if a:
            row = angle_table.add_row()
            row.cells[0].text = angle_name
            row.cells[1].text = a.get('iau_constellation', a.get('standard_constellation', '?'))
            row.cells[2].text = f'{a.get("iau_degree", a.get("standard_degree", 0)):.2f}°'
            row.cells[3].text = f'{a.get("standard_constellation", "?")} {a.get("standard_degree", 0):.2f}°'
            
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)


# ═══════════════════════════════════════════════════════════════════════════════
# AI CONTENT PARSER
# ═══════════════════════════════════════════════════════════════════════════════

def parse_ai_content(doc: Document, content: str):
    """
    Parse AI-generated markdown-ish content into styled docx paragraphs.
    Handles: ## headers, ### sub-headers, **bold**, *italic*, plain text,
    | tables |, and --- separators.
    """
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Section heading (##)
        if line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            p = doc.add_paragraph(style='SectionHeading')
            p.add_run(text)
            i += 1
            continue
        
        # Sub-section heading (###)
        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_paragraph(style='SubSectionHeading')
            p.add_run(text)
            i += 1
            continue
        
        # Chapter heading (#) — shouldn't normally appear in content, but handle it
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            p = doc.add_paragraph(style='ChapterHeading')
            p.add_run(text)
            i += 1
            continue
        
        # Horizontal rule
        if line.startswith('---') or line.startswith('***') or line.startswith('___'):
            # Add a thin line paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            # Use border instead of characters
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:color'), '999999')
            bottom.set(qn('w:space'), '1')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue
        
        # Table detection (| ... | ... |)
        if line.startswith('|') and '|' in line[1:]:
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tl = lines[i].strip()
                # Skip separator rows (|---|---|)
                if not re.match(r'^\|[\s\-:]+\|', tl.replace(' ', '')):
                    table_lines.append(tl)
                i += 1
            
            if table_lines:
                # Parse cells
                rows_data = []
                for tl in table_lines:
                    cells = [c.strip() for c in tl.split('|')[1:-1]]
                    rows_data.append(cells)
                
                if rows_data:
                    ncols = max(len(r) for r in rows_data)
                    table = doc.add_table(rows=0, cols=ncols)
                    table.style = 'Light Shading Accent 1'
                    
                    for ri, row_data in enumerate(rows_data):
                        row = table.add_row()
                        for ci, cell_text in enumerate(row_data):
                            if ci < ncols:
                                row.cells[ci].text = cell_text
                                for paragraph in row.cells[ci].paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(9)
                                        if ri == 0:
                                            run.font.bold = True
                    
                    doc.add_paragraph('')  # spacing after table
            continue
        
        # Regular paragraph with inline formatting
        p = doc.add_paragraph()
        parse_inline_formatting(p, line)
        i += 1


def parse_inline_formatting(paragraph, text: str):
    """
    Parse inline markdown formatting: **bold**, *italic*, ***bold italic***.
    """
    # Pattern to match **bold**, *italic*, and plain text
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'
    
    for match in re.finditer(pattern, text):
        if match.group(2):  # ***bold italic***
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(3):  # **bold**
            run = paragraph.add_run(match.group(3))
            run.bold = True
        elif match.group(4):  # *italic*
            run = paragraph.add_run(match.group(4))
            run.italic = True
        elif match.group(5):  # plain text
            paragraph.add_run(match.group(5))


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER ASSEMBLY
# ═══════════════════════════════════════════════════════════════════════════════

def add_chapter(doc: Document, title: str, content: str):
    """Add a chapter to the document with title and AI-generated content."""
    # Chapter heading
    p = doc.add_paragraph(style='ChapterHeading')
    p.add_run(title)
    
    # Parse and add content
    parse_ai_content(doc, content)


# ═══════════════════════════════════════════════════════════════════════════════
# CLOSING PAGE
# ═══════════════════════════════════════════════════════════════════════════════

def add_closing_page(doc: Document, chart: dict):
    """Add the closing page with computation notes."""
    add_page_break(doc)
    
    for _ in range(6):
        doc.add_paragraph('')
    
    p = doc.add_paragraph(style='Epigraph')
    p.add_run('⛎')
    
    p = doc.add_paragraph(style='Epigraph')
    p.add_run('"The stars don\'t lie. The systems that read them do."')
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # Technical notes
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'Chart calculated using Swiss Ephemeris (pyswisseph) with Moshier analytical + SE file ephemerides. '
        'IAU constellation boundaries (1930 epoch, precession-corrected). '
        'Whole Sign house system. '
        'Uranian hypothetical planets per Hamburg School ephemeris. '
        'Arabic Parts calculated using Hellenistic day/night chart formulas.'
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Generated by Stellaris-13 · {datetime.now().strftime("%B %d, %Y")}')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN GENERATOR
# ═══════════════════════════════════════════════════════════════════════════════

def generate_blueprint(
    chart: dict,
    birth_data: dict,
    ai_caller: Callable[[str], str],
    output_path: str,
    progress_callback: Optional[Callable[[int, str], None]] = None,
    max_tokens_per_chapter: int = 8000
) -> str:
    """
    Generate a complete celestial blueprint document.
    
    Args:
        chart: The computed chart data from engine.compute_chart()
        birth_data: Dict with 'date', 'time', 'location', 'lat', 'lon'
        ai_caller: Function that takes a prompt string and returns AI response text.
                   Signature: ai_caller(prompt: str) -> str
                   The caller should handle its own API config (model, max_tokens, etc.)
        output_path: Path to save the .docx file
        progress_callback: Optional callback(step_number, step_description) for UI progress
        max_tokens_per_chapter: Hint for the AI caller (not enforced here)
    
    Returns:
        Path to the generated document
    """
    
    def progress(step: int, desc: str):
        if progress_callback:
            progress_callback(step, desc)
        logger.info(f"Blueprint step {step}: {desc}")
    
    progress(0, "Initializing document...")
    
    # Create document
    doc = Document()
    setup_document_styles(doc)
    
    # Set page size to US Letter
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # Title page
    progress(1, "Creating title page...")
    create_title_page(doc, chart, birth_data)
    
    # Positions table
    progress(2, "Adding positions reference table...")
    add_positions_table(doc, chart)
    
    # Get chapter prompts
    chapters = get_chapter_prompts(chart)
    total_chapters = len(chapters)
    
    # Generate each chapter via AI
    for i, (title, prompt) in enumerate(chapters):
        step_num = 3 + i
        progress(step_num, f"Writing {title}...")
        
        try:
            # Build the full prompt with persona
            full_prompt = f"{ASTROLOGER_PERSONA}\n\n{prompt}"
            
            # Call AI — the caller handles model selection, temperature, etc.
            content = ai_caller(full_prompt)
            
            if not content or len(content.strip()) < 100:
                logger.warning(f"Short/empty response for {title}, adding placeholder")
                content = f"[Content generation incomplete for {title}. Re-run generation or add manually.]"
            
            # Add to document
            add_chapter(doc, title, content)
            
        except Exception as e:
            logger.error(f"Error generating {title}: {e}")
            add_chapter(doc, title, f"[Content generation failed: {e}]")
    
    # Closing page
    progress(3 + total_chapters, "Adding closing page...")
    add_closing_page(doc, chart)
    
    # Save document
    progress(4 + total_chapters, "Saving document...")
    doc.save(output_path)
    
    progress(5 + total_chapters, "Blueprint complete!")
    
    return output_path


# ═══════════════════════════════════════════════════════════════════════════════
# CONVENIENCE: Generate with Anthropic API
# ═══════════════════════════════════════════════════════════════════════════════

def make_anthropic_caller(api_key: str, model: str = "claude-sonnet-4-20250514", max_tokens: int = 8000):
    """
    Create an ai_caller function that uses the Anthropic API.
    
    Usage:
        caller = make_anthropic_caller("sk-ant-...")
        generate_blueprint(chart, birth_data, caller, "output.docx")
    """
    import anthropic
    
    client = anthropic.Anthropic(api_key=api_key)
    
    def caller(prompt: str) -> str:
        response = client.messages.create(
            model=model,
            max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}]
        )
        return response.content[0].text
    
    return caller


def make_openai_caller(api_key: str, model: str = "gpt-4o", max_tokens: int = 8000):
    """
    Create an ai_caller function that uses the OpenAI API.
    """
    from openai import OpenAI
    
    client = OpenAI(api_key=api_key)
    
    def caller(prompt: str) -> str:
        response = client.chat.completions.create(
            model=model,
            max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content
    
    return caller


def make_ollama_caller(model: str = "qwen2.5-coder:32b", base_url: str = "http://localhost:11434"):
    """
    Create an ai_caller function that uses a local Ollama model.
    Good for keeping everything on the Forge.
    """
    import requests
    
    def caller(prompt: str) -> str:
        response = requests.post(
            f"{base_url}/api/generate",
            json={
                "model": model,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "num_predict": 8000,
                    "temperature": 0.7
                }
            },
            timeout=300
        )
        response.raise_for_status()
        return response.json().get('response', '')
    
    return caller


# ═══════════════════════════════════════════════════════════════════════════════
# TESTING
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("Stellaris-13 Celestial Blueprint Generator v2")
    print("=" * 50)
    print(f"Chapters: {len(get_chapter_prompts({'placements': {}, 'angles': {}, 'aspects': []}))}")
    print("Ready. Use generate_blueprint() to create documents.")
    print()
    print("Quick test:")
    print("  from engine import compute_chart")
    print("  chart = compute_chart(1985, 12, 12, 10, 47, 0, -6, 41.5868, -93.625)")
    print("  caller = make_ollama_caller('qwen2.5-coder:32b')")
    print("  generate_blueprint(chart, {...}, caller, 'blueprint.docx')")
